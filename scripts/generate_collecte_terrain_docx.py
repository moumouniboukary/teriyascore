"""Génère TeriyaScore_Collecte_Donnees_Terrain.docx — guide de collecte pour le NeoScore."""
from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parents[1] / "docs" / "TeriyaScore_Collecte_Donnees_Terrain.docx"


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = table.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = val
    doc.add_paragraph()


def main() -> None:
    doc = Document()
    title = doc.add_heading("TeriyaScore — Guide de collecte des données terrain", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(
        "Solvabilité (NeoScore) · Secteur informel · Burkina Faso\n"
        f"Version 1.1 — {date.today().strftime('%d/%m/%Y')}"
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x4A, 0x6B, 0x5C)

    doc.add_paragraph(
        "Ce document recense toutes les informations à recueillir pour calculer, "
        "calibrer et améliorer le score de solvabilité TeriyaScore. "
        "Il couvre : l’application mobile, l’enquête KoboCollect, la visite agent terrain, "
        "le suivi crédit IMF et les labels pour le modèle ML."
    )

    doc.add_heading("1. Objectif de la collecte", level=1)
    doc.add_paragraph(
        "Le NeoScore (0–100) mesure la solvabilité alternative d’un travailleur informel "
        "sans bulletin de salaire ni historique bancaire. Seuil d’éligibilité crédit : 50/100."
    )
    add_table(
        doc,
        ["Priorité", "Type de donnée", "Usage"],
        [
            ["1 — Critique", "Outcomes crédit (remboursé / défaut)", "Entraînement et validation du ML"],
            ["2 — Élevée", "Cahier numérique (30 jours min.)", "Score basé sur l’activité réelle"],
            ["3 — Élevée", "Profil activité (app ou Kobo)", "Features de base du modèle"],
            ["4 — Moyenne", "Visite agent terrain", "Vérification et variables complémentaires"],
            ["5 — Moyenne", "Décision IMF (montant accordé / refus)", "Calibrage partenaire"],
        ],
    )

    doc.add_heading("2. Identité et contact", level=1)
    add_table(
        doc,
        ["Champ", "Format", "Obligatoire", "Source", "Remarques"],
        [
            ["Nom affiché", "Texte (1–120 car.)", "Oui", "App / Kobo", "Prénom et nom"],
            ["Téléphone", "+226 XX XX XX XX", "Oui", "App / Kobo", "Identifiant unique compte"],
            ["Genre", "homme | femme", "Recommandé", "Kobo / profil", "Cible femmes > 50 %"],
            ["Langue préférée", "fr | mr | dl | ff", "Oui", "App / Kobo", "Interface et assistance vocale"],
            ["Ville", "Texte", "Recommandé", "App onboarding", "Ex. Ouagadougou"],
            ["Zone / quartier", "Texte", "Recommandé", "App onboarding", "Localisation activité"],
            ["ID enquête / Kobo", "UUID", "Si Kobo", "KoboCollect", "submissionId pour traçabilité"],
        ],
    )

    doc.add_heading("3. Profil d’activité (NeoScore — features de base)", level=1)
    doc.add_paragraph(
        "Ces champs alimentent directement le calcul NeoScore (packages/neoscore + features.ts). "
        "Collecte via onboarding app ou formulaire KoboCollect (TeriyaScore_KoboCollect_XLSForm.xlsx)."
    )
    add_table(
        doc,
        ["Champ", "Valeurs possibles", "Obligatoire", "Impact score"],
        [
            ["Métier (metier)", "commerce, mecanique, artisanat, menuiserie, restauration, transport, agriculture, services", "Oui", "Segmentation métier"],
            ["Ancienneté activité (anciennete)", "m1 (<1 an), 1_2, 3_5, 6_10, p10 (>10 ans)", "Oui", "Régularité, croissance"],
            ["CA journalier estimé (ca_jour)", "m5k, 5_15k, 15_30k, 30_60k, 60_100k, p100k (FCFA)", "Oui", "Volume d’activité"],
            ["Participation tontine (tontine)", "oui | non", "Oui", "Régularité (+15 si oui)"],
            ["Cotisation tontine (tontine_cotis)", "Montant FCFA / mois", "Si tontine=oui", "Croissance (proxy ancienneté)"],
            ["Usage mobile money (mobile_money)", "jamais, occasionnel, regulier, quotidien", "Oui", "Régularité"],
            ["Statut compte bancaire (compte)", "non, oui_dormant, oui_actif", "Oui", "Gestion créances (+5 si actif)"],
            ["Charges fixes mensuelles (chargesFixesMensuelles)", "FCFA (loyer, famille…)", "Recommandé", "Capacité remboursement"],
            ["Saisonnalité (saisonnalite)", "stable, moderee, forte", "Recommandé", "Ajustement revenu estimé"],
            ["Garantie solidaire (garantieSolidaire)", "oui | non", "Optionnel", "Croissance (+8), relève segment"],
            ["Smartphone (telephone)", "aucun, basique, smartphone", "Recommandé", "Volume (+5 si smartphone)"],
            ["Impayés déclarés (impayes)", "0, m5k, 5_15k, 15_50k, p50k (FCFA dus)", "Recommandé", "Cross-check avec cahier"],
        ],
    )

    doc.add_heading("4. Enquête KoboCollect — champs complémentaires (recherche / calibrage)", level=1)
    doc.add_paragraph(
        "Champs du modèle POESAM 2026 et import API Kobo. Utiles pour recalibrer les poids "
        "avant d’avoir assez de labels crédit réels."
    )
    add_table(
        doc,
        ["Champ", "Valeurs possibles", "Usage"],
        [
            ["Âge (age)", "m25, 25_34, 35_44, 45_54, 55p", "Segmentation démographique"],
            ["Instruction (instruction)", "aucun, alpha, primaire, secondaire, superieur", "Accessibilité app"],
            ["Nb transactions / jour (nb_transactions)", "Entier", "Cross-check régularité"],
            ["Part ventes à crédit (part_credit)", "0_10 %, 10_25 %, 25_50 %, >50 %", "Feature partCredit"],
            ["Historique crédit déclaré (credit_hist)", "jamais, refuse, accorde", "Feature creditHist"],
            ["Besoin crédit (besoin_credit)", "m50k, 50_150k, 150_500k, 500k_2m, p2m", "Intention, pas le score"],
            ["Intérêt TeriyaScore (interet)", "Likert 1–5", "Recherche adoption (POESAM)"],
            ["Consentement (consentement)", "oui_libre, oui_anonyme, oui_benefice, non", "RGPD / éthique"],
        ],
    )

    doc.add_heading("5. Cahier numérique — activité réelle (30 derniers jours)", level=1)
    doc.add_paragraph(
        "Source principale pour un score crédible. Minimum recommandé : 4 semaines d’usage régulier "
        "avant décision crédit. Fenêtre d’analyse NeoScore : 30 jours glissants."
    )
    add_table(
        doc,
        ["Type opération", "Champs à enregistrer", "Impact NeoScore"],
        [
            ["Vente (vente)", "montant FCFA, date, libellé, client (optionnel), canal (espèces/mobile money)", "salesLast30Fcfa, volume"],
            ["Créance (creance)", "montant, client, date échéance, statut (ouverte/en_retard/réglée)", "openDebtsFcfa, overdueDebtsCount, dettes"],
            ["Dépense (depense)", "montant, catégorie, date", "expensesLast30Fcfa, marge nette"],
            ["Stock (stock)", "entrée/sortie, article, quantité, prix unitaire", "Cohérence activité"],
        ],
    )
    add_table(
        doc,
        ["Indicateur calculé", "Formule / source", "Critère NeoScore"],
        [
            ["opsLast30Days", "Nombre total d’opérations sur 30 j", "Régularité, croissance"],
            ["salesLast30Fcfa", "Somme ventes 30 j", "Volume"],
            ["openDebtsFcfa", "Somme créances non réglées", "Gestion créances (−)"],
            ["overdueDebtsCount", "Nb créances en_retard", "Gestion créances (−)"],
            ["partCredit", "Ratio créances / (ventes + créances)", "Volume"],
            ["impayes", "Cap à 4 retards", "Gestion créances (−)"],
            ["creditHist", "0 aucun, 1 créances réglées, 2 crédit approuvé", "Croissance"],
            ["expensesLast30Fcfa", "Somme dépenses 30 j", "Volume (marge nette)"],
            ["activeWeeksLast30", "Semaines avec ≥1 opération", "Éligibilité (min. 4 sem.)"],
            ["salesVsDeclaredRatio", "Ventes 30 j / (CA déclaré × 30)", "Cohérence déclaratif / cahier"],
            ["tontineCotisations30Fcfa", "Module tontine (30 j)", "Régularité, capacité remboursement"],
            ["monthlyFixedChargesFcfa", "Profil charges fixes", "Capacité remboursement"],
        ],
    )
    doc.add_paragraph(
        "Règles d’éligibilité crédit (v1.1) : score ≥ 50 ET activité min. "
        "(5 opérations ou 4 semaines actives sur 30 j) ET capacité remboursement "
        "(mensualité ≤ 35 % du revenu mensuel net estimé, offre plafonnée)."
    )
    doc.add_paragraph("Clients informels : nom, téléphone (optionnel), note — pour lier les créances.")

    doc.add_heading("6. Visite agent terrain (phase pilote / IMF)", level=1)
    doc.add_paragraph(
        "Données observées sur place — les champs charges fixes, saisonnalité et "
        "garantie solidaire sont désormais intégrés au profil app et au NeoScore."
    )
    add_table(
        doc,
        ["Champ", "Type", "Obligatoire", "Remarques"],
        [
            ["Date visite", "Date", "Oui", ""],
            ["Agent / enquêteur", "Texte", "Oui", "Traçabilité"],
            ["Activité visible cohérente", "oui | non | partiel", "Oui", "Stock, clients, outillage"],
            ["Photo lieu d’activité", "Image", "Recommandé", "Preuve terrain"],
            ["Estimation CA observé / jour", "FCFA", "Recommandé", "Comparer au CA déclaré"],
            ["Saisonnalité", "Texte", "Optionnel", "Agriculture, fêtes"],
            ["Charges fixes mensuelles", "FCFA (loyer, famille…)", "Recommandé", "Capacité remboursement"],
            ["Groupe solidaire / garantie", "oui | non + détail", "Optionnel", "Pratique IMF"],
            ["Réputation locale", "Likert 1–5", "Optionnel", "Volonté de payer"],
            ["Commentaire agent", "Texte libre", "Optionnel", ""],
        ],
    )

    doc.add_heading("7. Suivi crédit et labels ML (priorité n°1)", level=1)
    doc.add_paragraph(
        "À chaque demande de crédit, un snapshot des features est figé (featuresSnapshot). "
        "La clôture avec outcome alimente l’entraînement ML."
    )
    add_table(
        doc,
        ["Champ", "Valeurs", "Quand", "Responsable"],
        [
            ["reference", "Réf. unique demande", "Soumission", "Système"],
            ["montantDemandeFcfa", "FCFA", "Soumission", "Utilisateur"],
            ["usage", "equipement | stock | urgence | autre", "Soumission", "Utilisateur"],
            ["modaliteRemboursement", "Texte", "Soumission", "Utilisateur / IMF"],
            ["statut", "soumise → en_examen → approuvee/refusee → decaissee → cloturee", "Cycle crédit", "IMF / admin"],
            ["dateDecaissement", "Date", "Décaissement", "IMF"],
            ["dateEcheance", "Date", "Décaissement", "IMF"],
            ["dateCloture", "Date", "Fin crédit", "IMF / admin"],
            ["outcome", "rembourse_ok | defaut", "Clôture", "IMF / admin — LABEL ML"],
            ["motifDecision", "Texte", "Refus / défaut", "IMF"],
            ["montantAccordeFcfa", "FCFA", "Décision IMF", "IMF (recommandé)"],
            ["mensualiteFcfa", "FCFA", "Décaissement", "Calcul capacité remboursement"],
        ],
    )
    doc.add_paragraph(
        "Objectif calibrage ML : minimum 200–500 crédits clôturés avec outcome connu "
        "(mix rembourse_ok et defaut)."
    )

    doc.add_heading("8. Consentements (obligatoire avant partage IMF)", level=1)
    add_table(
        doc,
        ["Type consentement", "Description", "Obligatoire crédit IMF"],
        [
            ["anonymisation_recherche", "Données anonymisées pour recherche / ML", "Non"],
            ["partage_imf", "Partage profil et NeoScore avec IMF partenaire", "Oui pour crédit"],
            ["marketing_partenaires", "Communications partenaires", "Non"],
        ],
    )

    doc.add_heading("9. Résultats NeoScore (générés — ne pas collecter manuellement)", level=1)
    add_table(
        doc,
        ["Indicateur", "Échelle", "Seuil / règle"],
        [
            ["Score global (valeur)", "0–100", "Éligible si ≥ 50 + activité + capacité"],
            ["Segment", "A, B, C, D", "A=régulier stable, D=exclusion"],
            ["Critère régularité", "0–100", "Poids 30 %"],
            ["Critère volume", "0–100", "Poids 25 % (incl. marge nette)"],
            ["Critère gestion créances", "0–100", "Poids 25 %"],
            ["Critère croissance", "0–100", "Poids 20 %"],
            ["Capacité remboursement", "max 35 % revenu mensuel net", "Plafond offre crédit"],
            ["Offre crédit (si éligible)", "min 50k, max min(score×3k, capacité) FCFA", "Validité 7 jours"],
            ["dataQuality.warnings", "Liste alertes", "Écart CA, activité faible, etc."],
        ],
    )

    doc.add_heading("10. Checklist pilote Ouagadougou", level=1)
    checklist = [
        "□ 100–200 enquêtes KoboCollect complétées (profil + consentement)",
        "□ 50+ utilisateurs avec cahier actif ≥ 4 semaines",
        "□ Chaque demande crédit : featuresSnapshot + dates décaissement/échéance",
        "□ Clôture systématique : outcome rembourse_ok ou defaut",
        "□ 10+ visites agent terrain avec photo et commentaire",
        "□ Export CSV Kobo → import API POST /kobo/import",
        "□ Ré-entraînement ML après 200 labels : POST /admin/ml/retrain",
    ]
    for item in checklist:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("11. Contacts et fichiers de référence", level=1)
    doc.add_paragraph("Formulaire Kobo : DAMINA&POESAM_2026/TeriyaScore_KoboCollect_XLSForm.xlsx")
    doc.add_paragraph("Modèle recherche : DAMINA&POESAM_2026/TeriyaScore_NeoScore_Model.py")
    doc.add_paragraph("Documentation ML : docs/ml-scoring.md")
    doc.add_paragraph("API import Kobo : POST /kobo/import (soumissions JSON)")
    doc.add_paragraph("Clôture solvabilité : POST /partners/applications/:id/outcome")

    doc.add_paragraph()
    p = doc.add_paragraph("Document généré pour TeriyaScore — usage interne équipe terrain, IMF et recherche.")
    p.runs[0].italic = True

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    main()
